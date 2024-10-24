from flask import Flask, render_template, request, send_file
import os
import fitz  # PyMuPDF
from docx import Document
from PyPDF2 import PdfFileReader, PdfFileWriter
import pdfkit
from docx import Document

app = Flask(__name__)

# Set the upload folder
UPLOAD_FOLDER = 'uploads/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Ensure the uploads folder exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


# Home route to render the HTML form
@app.route('/')
def home():
    return render_template('index.html')


# Route to handle file upload and conversion
@app.route('/convert', methods=['POST'])
def convert_file():
    if 'pdf_file' in request.files:
        file = request.files['pdf_file']
        if file.filename.endswith('.pdf'):
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(pdf_path)

            docx_path = os.path.splitext(pdf_path)[0] + ".docx"
            try:
                pdf_to_docx(pdf_path, docx_path)
            except Exception as e:
                return f"An error occurred during conversion: {e}"

            return send_file(docx_path, as_attachment=True)

    elif 'word_file' in request.files:
        file = request.files['word_file']
        if file.filename.endswith('.docx'):
            word_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(word_path)

            pdf_path = os.path.splitext(word_path)[0] + ".pdf"
            try:
                docx_to_pdf(word_path, pdf_path)
            except Exception as e:
                return f"An error occurred during conversion: {e}"

            return send_file(pdf_path, as_attachment=True)

    return "Invalid file format. Please upload a PDF or Word file."


def pdf_to_docx(pdf_path, docx_path):
    pdf_document = fitz.open(pdf_path)
    docx_document = Document()
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text = page.get_text()
        if text:
            docx_document.add_paragraph(text)
    docx_document.save(docx_path)


def docx_to_pdf(docx_path, pdf_path):
    # Read the Word document
    docx_document = Document(docx_path)
    
    # Create a temporary HTML file to hold the document contents
    html_content = ""
    for para in docx_document.paragraphs:
        html_content += f"<p>{para.text}</p>\n"
    
    # Save the HTML content to a temporary HTML file
    with open("temp_file.html", "w") as html_file:
        html_file.write(html_content)
    
    # Convert the HTML file to PDF using pdfkit
    pdfkit.from_file("temp_file.html", pdf_path)


if __name__ == '__main__':
    app.run(debug=True)
